from docx import Document
import sys

doc = Document(sys.argv[1])
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
print("PARAGRAPHS")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style is not None else "[none]"
    print(f"{i:03d}\t{style}\t{p.text[:220]}")
print("TABLES")
for i, table in enumerate(doc.tables):
    first = " || ".join(cell.text.replace("\n", " / ")[:120] for cell in table.rows[0].cells)
    print(f"T{i:02d}\t{len(table.rows)}x{len(table.columns)}\t{first}")
